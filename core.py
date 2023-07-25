import os
import re

import docx
import xlrd
from PyQt5.QtWidgets import QMessageBox
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm

print("错题集生成器", end='\n')
print("作者：Leonard", end='\n')

split_pattern=r',|，| '

def generate(jyfoldername,cuotifkbname,classes,subject,weeknumber,date,widget):
    textBrowser = widget.textBrowser
    jyfoldername=jyfoldername+'\\'
    excel = xlrd.open_workbook(cuotifkbname, encoding_override="utf-8")
    sheet = excel.sheets()[0]
    NAMES = sheet.col_values(0)
    ctdoc = ['']

    #备份
    import shutil
    if os.path.exists("backup\\"):
        shutil.rmtree("backup\\")
    if os.path.exists("学生\\"):
        shutil.copytree("学生\\", "backup\\")

    # 得到名字name和对应的错题号problemnumber（PN）
    for i in range(1, NAMES.__len__()):
        # 得到pn
        stupath = "学生\\" + sheet.cell_value(i, 0) + "\\" + subject + "\\" + "第" + weeknumber + "周错题集.docx"
        if not os.path.exists("学生\\" + sheet.cell_value(i, 0) + "\\" + subject + "\\"):
            os.makedirs("学生\\" + sheet.cell_value(i, 0) + "\\" + subject + "\\")
        if not os.path.exists(stupath):
            doc = docx.Document()
            doc.sections[0].top_margin = Cm(1.27)
            doc.sections[0].bottom_margin = Cm(1.27)
            doc.sections[0].left_margin = Cm(1.27)
            doc.sections[0].right_margin = Cm(1.27)

            h = doc.sections[0].header
            h.paragraphs[0].text = classes + "    " + sheet.cell_value(i, 0) + "    " + subject + "    第" + weeknumber + "周"
            p = h.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            doc.save(stupath)

        ctdoc.append(Document(stupath))
        PN = str(sheet.cell_value(i, 1))

        # 添加日期标题
        ctdoc[i].add_heading(date, level=1)

        # 没带的话
        if (PN == "未带") or (PN == "未交"):
            for root, dirs, files in os.walk(jyfoldername, topdown=False):
                for f in files:
                    if f.endswith(".png") and f[:-4].isdigit():
                        ctdoc[i].add_picture(jyfoldername + f,width=Cm(18))  # 直接操作

            textBrowser.append(sheet.cell_value(i, 0) + str(date) + "日本科目未带")
            # ctdoc.add_paragraph("未交")
            pass

        # 全对的话
        elif PN == "全对":
            textBrowser.append(sheet.cell_value(i, 0) + str(date) + "日本科目全对")
            ctdoc[i].add_paragraph("全对")
            pass

        # 其他情况
        else:
            pn = re.split(split_pattern, PN)
            for b in pn:
                if b == "":
                    continue
                c = int(float(b))
                if c < 10:
                    a = '0' + str(c)
                else:
                    a = str(c)
                try:
                    ctdoc[i].add_picture(jyfoldername + a + '.png',width=Cm(18))
                except FileNotFoundError:
                    textBrowser.append("<font color='red'>错误：</font>检查你的错题反馈表是否输入了讲义文件夹中没有的题目")
                    textBrowser.append(f"应为{sheet.cell_value(i, 0)}错的第{a}题是否出现在右侧列表？\n")
                    if QMessageBox.warning(widget,'错误','忽略并跳过本题？',QMessageBox.Yes|QMessageBox.No,QMessageBox.Yes)==QMessageBox.Yes:
                        continue
                    else:
                        raise Exception("你的错题反馈表含有讲义文件夹中没有的题目")
            textBrowser.append(sheet.cell_value(i, 0) + str(date) + '日错题集已完成')

    textBrowser.append("<font color='blue'>开始保存错题集...</font>")
    for i in range(1, NAMES.__len__()):
        stupath = "学生\\" + sheet.cell_value(i, 0) + "\\" + subject + "\\" + "第" + weeknumber + "周错题集.docx"
        ctdoc[i].save(stupath)
        textBrowser.append(sheet.cell_value(i, 0) + '的' + subject + '错题集已保存')

    textBrowser.append("<font color='blue'>本次作业人数：" + str(NAMES.__len__()) + "</font>\n")

    textBrowser.append("正在保存记录。。。")

    if not os.path.exists('记录.docx'):
        record = Document()
        record.save('记录.docx')

    record = Document('记录.docx')
    record.add_heading(date, level=1)
    record.add_paragraph(jyfoldername+cuotifkbname)
    record.add_paragraph(subject)
    record.add_paragraph("\n")
    record.save('记录.docx')

