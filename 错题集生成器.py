import re
from docx import *
import os
import docx.shared
import xlrd

print("错题集生成器", end='\n')
print("作者：Leonard", end='\n')
path = ".\\"
jydocname = input("讲义文件名:") + ".docx"  # 原讲义提取出题目并标括号
cuotifkbname = input("错题反馈表文件名:") + ".xlsx"
pattern1 = "(" + input("错题日期:") + ")"

l = [];
row = ""
jydoc = Document(path + jydocname);
for i in jydoc.paragraphs:

    if (re.findall(pattern1, i.text)).__len__() != 0:

        l.append(row)
        row = "";
    else:
       pass
    row += i.text + "\n";
# 想办法对row加入图片
l.append(row)

excel = xlrd.open_workbook(cuotifkbname, encoding_override="utf-8")
sheet = excel.sheets()[0]

NAMES = sheet.col_values(0)

# 得到名字name和对应的错题号problemnumber（PN）
for i in range(1, NAMES.__len__()):
    # 得到pn
    stupath = path + "学生\\" + sheet.cell_value(i, 0) + "\\"
    if not os.path.exists(stupath):
        doc = docx.Document()
        os.mkdir(stupath)
        doc.save(stupath + '错题集.docx')

    ctdoc = Document(stupath + '错题集.docx')
    PN=str(sheet.cell_value(i, 1))
    if (PN == "未带") or (PN == "未交"):
        pn=list(range(1,l.__len__()))
    elif (PN == "全对"):
        pn=[]
    else:
        pn = PN.split("，")

    for b in pn:
        if (b == ""):
            continue
        a = int(float(b))
        ctdoc.add_paragraph(l[int(a)])
        ctdoc.add_paragraph(f"\n")

        print(l[int(a)])
        print(f'\n')

    ctdoc.save(stupath + '错题集.docx')

print("本次作业人数：" + str(NAMES.__len__()) + "\n")
print("分别为：\n")
for s in NAMES:  # CR!
    print(s, end="\n")

print("正在保存。。。")

if not os.path.exists(path + '记录.docx'):
    record = Document()
    record.save(path + '记录.docx')

record=Document(path + '记录.docx')
record.add_paragraph(pattern1)
record.add_paragraph(jydocname)
record.add_paragraph(cuotifkbname)
record.add_paragraph("\n")
record.save(path + '记录.docx')

print("工作完成！！！！")

input()
