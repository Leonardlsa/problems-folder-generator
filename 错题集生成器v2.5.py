import os
import docx
import xlrd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm

print("错题集生成器", end='\n')
print("作者：Leonard", end='\n')
path = ".\\"
jyfoldername = path + input("讲义文件夹名:") + '\\'  # 原讲义提取出题目并标括号
cuotifkbname = input("错题反馈表文件名:") + ".xlsx"
subject= input("本学科名：")

date = input("日期：")
weeknumber=input("此为第几周：")

excel = xlrd.open_workbook(cuotifkbname, encoding_override="utf-8")
sheet = excel.sheets()[0]

NAMES = sheet.col_values(0)
ctdoc = ['']

# 得到名字name和对应的错题号problemnumber（PN）
for i in range(1, NAMES.__len__()):
    # 得到pn
    stupath = path + "学生\\" + sheet.cell_value(i, 0) + "\\" +subject+ "\\"+"第"+weeknumber+"周错题集.docx"
    if not os.path.exists(path + "学生\\"):
        os.mkdir(path + "学生\\")
    if not os.path.exists(path + "学生\\" + sheet.cell_value(i, 0) + "\\"):
        os.mkdir(path + "学生\\"+ sheet.cell_value(i, 0) + "\\")
    if not os.path.exists(path + "学生\\" + sheet.cell_value(i, 0) + "\\"+ subject + "\\"):
        os.mkdir(path + "学生\\" + sheet.cell_value(i, 0) + "\\" + subject + "\\")
    if not os.path.exists(stupath):
        doc = docx.Document()
        h=doc.sections[0].header
        h.paragraphs[0].text = "天任教育    A4    "+sheet.cell_value(i, 0)+"    "+subject+"    第"+weeknumber+"周"
        p=h.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.sections[0].top_margin = Cm(1.27)
        doc.sections[0].bottom_margin = Cm(1.27)
        doc.sections[0].left_margin = Cm(1.27)
        doc.sections[0].right_margin = Cm(1.27)
        doc.save(stupath)

    ctdoc.append(Document(stupath))
    PN = str(sheet.cell_value(i, 1))

    # 添加日期标题
    ctdoc[i].add_heading(date, level=1)

    # 没带的话
    if (PN == "未带") or (PN == "未交"):
        for root, dirs, files in os.walk(jyfoldername, topdown=False):
             for f in files:
                 ctdoc[i].add_picture(jyfoldername + f)  # 直接操作

        print(sheet.cell_value(i, 0) + str(date) + "日本科目未带")
         # ctdoc.add_paragraph("未交")
        pass

    # 全对的话
    elif PN == "全对":
        print(sheet.cell_value(i, 0) + str(date) + "日本科目全对")
        ctdoc[i].add_paragraph("全对")
        pass

    # 其他情况
    else:
        pn = PN.split("，")
        for b in pn:
            if b == "":
                continue
            c = int(float(b))
            if c < 10:
                a = '0' + str(c)
            else:
                a = str(c)
            ctdoc[i].add_picture(jyfoldername + a + '.png')

    print(sheet.cell_value(i, 0) + str(date) + '日错题集已完成')

print('开始保存错题集...')
for i in range(1, NAMES.__len__()):
    stupath = path + "学生\\" + sheet.cell_value(i, 0) + "\\" +subject+ "\\"+"第"+weeknumber+"周错题集.docx"
    ctdoc[i].save(stupath)
    print(sheet.cell_value(i, 0) + '的'+subject+'错题集已保存')

print("本次作业人数：" + str(NAMES.__len__()) + "\n")

print("正在保存记录。。。")

if not os.path.exists(path + '记录.docx'):
    record = Document()
    record.save(path + '记录.docx')

record = Document(path + '记录.docx')
record.add_heading(date, level=1)
record.add_paragraph(jyfoldername)
record.add_paragraph(cuotifkbname)
record.add_paragraph("\n")
record.save(path + '记录.docx')

print("工作完成！！！！")

input("键入Enter键以关闭. . . . . . ")
