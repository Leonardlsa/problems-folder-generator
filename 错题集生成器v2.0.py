import os
import docx
import xlrd
from docx import Document

print("错题集生成器", end='\n')
print("作者：Leonard", end='\n')
path = ".\\"
jyfoldername = path + input("讲义文件夹名:") + '\\'  # 原讲义提取出题目并标括号
cuotifkbname = input("错题反馈表文件名:") + ".xlsx"
date = input("日期：")
l = [];

excel = xlrd.open_workbook(cuotifkbname, encoding_override="utf-8")
sheet = excel.sheets()[0]

NAMES = sheet.col_values(0)

# 得到名字name和对应的错题号problemnumber（PN）
for i in range(1, NAMES.__len__()):
    # 得到pn
    stupath = path + "学生\\" + sheet.cell_value(i, 0) + "\\"
    if not os.path.exists(stupath):
        doc = docx.Document()
        os.mkdir(stupath)
        doc.save(stupath + '错题集.docx')

    ctdoc = Document(stupath + '错题集.docx')
    PN = str(sheet.cell_value(i, 1))


    ctdoc.add_heading(date, level=1)

    if (PN == "未带") or (PN == "未交"):
        pn = list(range(1, l.__len__()))
        print(sheet.cell_value(i, 0) + str(date) + "日本科目未带")
        #ctdoc.add_paragraph("未交")

    elif PN == "全对":
        pn=[]
        print(sheet.cell_value(i, 0) + str(date) + "日本科目全对")
        ctdoc.add_paragraph("全对")

    else:
        pn = PN.split("，")

    for b in pn:
        if b == "":
            continue

        ctdoc.add_picture(jyfoldername + str(int(float(b))) + '.png')

    ctdoc.save(stupath + '错题集.docx')
    print(sheet.cell_value(i, 0) + str(date) + '日错题集已完成')

print("本次作业人数：" + str(NAMES.__len__()) + "\n")
print("分别为：\n")
for s in NAMES:  # CR!
    print(s, end="\n")

print("正在保存。。。")

if not os.path.exists(path + '记录.docx'):
    record = Document()
    record.save(path + '记录.docx')

record = Document(path + '记录.docx')
ctdoc.add_heading(date, level=1)
record.add_paragraph(jyfoldername)
record.add_paragraph(cuotifkbname)
record.add_paragraph("\n")
record.save(path + '记录.docx')

print("工作完成！！！！")

input()
